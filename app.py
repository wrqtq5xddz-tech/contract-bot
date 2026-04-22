import asyncio
import io
import os
import re
from datetime import datetime

import httpx
import json
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from dotenv import load_dotenv

load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "anthropic/claude-sonnet-4-5")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

SYSTEM_PROMPT = """Ты — опытный юрист-аналитик, специализирующийся на российском \
договорном праве и законодательстве о долевом строительстве (214-ФЗ).

Проанализируй договор и верни ТОЛЬКО валидный JSON без markdown и пояснений.

Уровни риска: HIGH — критическая угроза, MEDIUM — существенный недостаток, LOW — замечание.

Если договор не является ДДУ по 214-ФЗ — укажи is_ddu: false, overall_status: "NOT_APPLICABLE".

Формат ответа:
{
  "overall_risk_level": "HIGH|MEDIUM|LOW",
  "risks": [
    {
      "severity": "HIGH|MEDIUM|LOW",
      "category": "Категория риска",
      "title": "Краткое название",
      "clause": "п. X.X",
      "description": "Описание проблемы",
      "recommendation": "Рекомендация"
    }
  ],
  "fz214": {
    "is_ddu": true,
    "status": "COMPLIANT|VIOLATIONS_FOUND|NOT_APPLICABLE",
    "violations": ["Описание нарушения 1"]
  },
  "summary": "Краткое резюме 3-5 предложений"
}"""

SEVERITY_LABEL = {"HIGH": "Высокий", "MEDIUM": "Средний", "LOW": "Низкий"}
SEVERITY_COLOR = {
    "HIGH": RGBColor(0xC0, 0x00, 0x00),
    "MEDIUM": RGBColor(0xFF, 0x8C, 0x00),
    "LOW": RGBColor(0x00, 0x70, 0xC0),
}


# ── Парсинг файлов ─────────────────────────────────────────────────────────

def parse_file(content: bytes, file_name: str) -> str:
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext == "pdf":
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)

    if ext == "docx":
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(parts)

    if ext == "txt":
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue

    raise ValueError(f"Неподдерживаемый формат: .{ext}")


# ── Вспомогательный парсер ответа LLM ─────────────────────────────────────

def _parse_llm_json(raw: str) -> dict | list:
    if not raw or not raw.strip():
        raise ValueError("LLM вернул пустой ответ")
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    m = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", raw)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"[\[{][\s\S]+[\]}]", raw)
    if m:
        return json.loads(m.group(0))
    raise ValueError(f"Не удалось разобрать ответ LLM. Начало: {raw[:300]}")


# ── LLM: анализ рисков ─────────────────────────────────────────────────────

async def analyze(text: str) -> dict:
    truncated = text[:100_000]
    user_prompt = f"""Проанализируй договор на юридические риски.

КАТЕГОРИИ РИСКОВ:
1. Ответственность сторон — несбалансированная ответственность
2. Сроки и неустойки — нереалистичные сроки, завышенные/заниженные санкции
3. Условия расторжения — одностороннее расторжение, несправедливые условия
4. Гарантийные обязательства — отсутствие или короткие гарантийные сроки
5. Оплата — скрытые платежи, непрозрачные расчёты
6. Форс-мажор — слишком широкое/узкое определение
7. Споры и подсудность — невыгодная подсудность

ПРОВЕРКА НА 214-ФЗ (если ДДУ):
- Описание объекта строительства (ст. 4 ч. 4 п. 1)
- Цена и порядок оплаты (ст. 4 ч. 4 п. 3)
- Срок передачи объекта (ст. 4 ч. 4 п. 2)
- Гарантийный срок не менее 5 лет (ст. 7 ч. 5)
- Неустойка 1/300 → 1/150 ставки ЦБ за просрочку (ст. 6 ч. 2)
- Условия расторжения дольщиком (ст. 9)
- Гос. регистрация договора (ст. 17)

ТЕКСТ ДОГОВОРА:
---
{truncated}
---"""

    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://contract-analyzer.local",
                "X-Title": "Contract Legal Risk Analyzer",
            },
            json={
                "model": OPENROUTER_MODEL,
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0.1,
                "max_tokens": 4096,
            },
        )
        resp.raise_for_status()

    data = resp.json()
    raw = data["choices"][0]["message"]["content"]
    return _parse_llm_json(raw)


# ── LLM: правки для HIGH-рисков ────────────────────────────────────────────

async def correct_risks(contract_text: str, high_risks: list[dict]) -> list[dict]:
    risks_desc = "\n".join(
        f"- {r.get('title', '')} ({r.get('clause', '?')}): {r.get('description', '')}"
        for r in high_risks
    )
    user_prompt = f"""В договоре выявлены риски ВЫСОКОГО уровня. \
Для каждого риска предложи конкретную правку формулировки.

РИСКИ ВЫСОКОГО УРОВНЯ:
{risks_desc}

Для каждого риска верни объект:
- clause        — ссылка на пункт договора
- risk_title    — название риска
- original_excerpt — дословная цитата проблемного места из договора (до 300 символов, \
или пустая строка если пункт отсутствует)
- suggested_text — предлагаемая новая редакция этого пункта (готовый юридический текст)
- rationale     — почему правка снижает риск (1-2 предложения)

ТЕКСТ ДОГОВОРА:
---
{contract_text[:80_000]}
---

Верни ТОЛЬКО валидный JSON без markdown:
{{
  "corrections": [
    {{
      "clause": "п. X.X",
      "risk_title": "...",
      "original_excerpt": "...",
      "suggested_text": "...",
      "rationale": "..."
    }}
  ]
}}"""

    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://contract-analyzer.local",
                "X-Title": "Contract Legal Risk Analyzer",
            },
            json={
                "model": OPENROUTER_MODEL,
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "Ты — опытный юрист. Предлагай конкретные, юридически грамотные "
                            "правки формулировок договора на русском языке. "
                            "Верни ТОЛЬКО валидный JSON."
                        ),
                    },
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0.2,
                "max_tokens": 4096,
            },
        )
        resp.raise_for_status()

    data = resp.json()
    raw = data["choices"][0]["message"]["content"]
    parsed = _parse_llm_json(raw)
    if isinstance(parsed, dict):
        return parsed.get("corrections", [])
    return parsed  # на случай если LLM вернул массив напрямую


# ── Word: вспомогательные функции ──────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _bold_run(para, text: str, size: int = 11, color: RGBColor | None = None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ── Word: генерация справки ────────────────────────────────────────────────

def build_word_report(
    data: dict,
    doc_name: str = "",
    corrections: list[dict] | None = None,
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    # Заголовок
    t = doc.add_heading("СПРАВКА", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Результаты анализа договора на юридические риски")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # Мета
    meta = doc.add_paragraph()
    meta.add_run("Дата анализа: ").bold = True
    meta.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
    if doc_name:
        m2 = doc.add_paragraph()
        m2.add_run("Документ: ").bold = True
        m2.add_run(doc_name)
    doc.add_paragraph()

    # 1. Общий уровень риска
    level = data.get("overall_risk_level", "—")
    doc.add_heading("1. Общий уровень риска", level=1)
    p = doc.add_paragraph()
    _bold_run(p, SEVERITY_LABEL.get(level, level), size=14, color=SEVERITY_COLOR.get(level))
    doc.add_paragraph()

    # 2. Резюме
    doc.add_heading("2. Краткое резюме", level=1)
    doc.add_paragraph(data.get("summary", "—"))
    doc.add_paragraph()

    # 3. Таблица рисков
    risks = data.get("risks", [])
    doc.add_heading("3. Выявленные риски", level=1)
    if not risks:
        doc.add_paragraph("Существенных рисков не выявлено.")
    else:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        headers = ["№", "Уровень", "Категория / Пункт", "Описание", "Рекомендация"]
        col_widths = [Inches(0.3), Inches(0.9), Inches(1.4), Inches(2.2), Inches(2.0)]
        for cell, hdr, w in zip(hdr_cells, headers, col_widths):
            cell.width = w
            _set_cell_bg(cell, "1F4E79")
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

        for idx, r in enumerate(risks, 1):
            sev = r.get("severity", "")
            row_cells = tbl.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(9)

            sev_run = row_cells[1].paragraphs[0].add_run(SEVERITY_LABEL.get(sev, sev))
            sev_run.bold = True
            sev_run.font.size = Pt(9)
            if sev in SEVERITY_COLOR:
                sev_run.font.color.rgb = SEVERITY_COLOR[sev]

            clause = r.get("clause", "")
            cat_text = r.get("category", "") + (f"\n{clause}" if clause else "")
            row_cells[2].paragraphs[0].add_run(cat_text).font.size = Pt(9)

            tr = row_cells[3].paragraphs[0].add_run(r.get("title", "") + "\n")
            tr.bold = True
            tr.font.size = Pt(9)
            row_cells[3].paragraphs[0].add_run(r.get("description", "")).font.size = Pt(9)

            row_cells[4].paragraphs[0].add_run(r.get("recommendation", "")).font.size = Pt(9)

            bg = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
            for cell in row_cells:
                _set_cell_bg(cell, bg)

    doc.add_paragraph()

    # 4. 214-ФЗ
    fz = data.get("fz214", {})
    doc.add_heading("4. Проверка соответствия 214-ФЗ", level=1)
    if not fz.get("is_ddu", False):
        doc.add_paragraph("Договор не является ДДУ — проверка по 214-ФЗ не применима.")
    else:
        status = fz.get("status", "")
        status_map = {
            "COMPLIANT": ("Соответствует требованиям 214-ФЗ", RGBColor(0x37, 0x86, 0x1D)),
            "VIOLATIONS_FOUND": ("Обнаружены нарушения 214-ФЗ", RGBColor(0xC0, 0x00, 0x00)),
        }
        label, color = status_map.get(status, (status, None))
        _bold_run(doc.add_paragraph(), label, color=color)
        violations = fz.get("violations", [])
        if violations:
            doc.add_paragraph("Нарушения:")
            for v in violations:
                doc.add_paragraph(f"• {v}", style="List Bullet")

    doc.add_paragraph()

    # 5. Рекомендуемые правки (опционально)
    if corrections:
        doc.add_heading("5. Рекомендуемые правки по рискам высокого уровня", level=1)
        doc.add_paragraph(
            "Ниже приведены предлагаемые изменения формулировок договора "
            "для устранения рисков высокого уровня."
        )
        doc.add_paragraph()

        for i, c in enumerate(corrections, 1):
            clause = c.get("clause", "")
            h2 = doc.add_heading(
                f"5.{i}. {c.get('risk_title', f'Правка {i}')}" + (f" ({clause})" if clause else ""),
                level=2,
            )

            orig = c.get("original_excerpt", "").strip()
            if orig:
                doc.add_paragraph("Текущая редакция:").runs[0].bold = True
                orig_p = doc.add_paragraph(orig)
                orig_p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            doc.add_paragraph("Предлагаемая редакция:").runs[0].bold = True
            sug_p = doc.add_paragraph(c.get("suggested_text", ""))
            sug_p.runs[0].font.color.rgb = RGBColor(0x37, 0x86, 0x1D)

            rationale = c.get("rationale", "")
            if rationale:
                rat_p = doc.add_paragraph(f"Обоснование: {rationale}")
                rat_p.runs[0].italic = True
                rat_p.runs[0].font.size = Pt(9)

            doc.add_paragraph()

    # Оговорка
    disc = doc.add_paragraph(
        "Настоящая справка подготовлена с использованием технологий искусственного интеллекта "
        "и носит информационный характер. Она не является юридической консультацией и не может "
        "заменить заключение квалифицированного юриста."
    )
    disc.runs[0].italic = True
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── UI: отображение результатов ────────────────────────────────────────────

def show_results(data: dict, doc_name: str = "", corrections: list[dict] | None = None):
    level = data.get("overall_risk_level", "—")
    level_color = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(level, "⚪")

    st.markdown(f"## {level_color} Общий уровень риска: **{level}**")
    st.info(data.get("summary", ""))

    risks = data.get("risks", [])
    if risks:
        st.markdown("### ⚠️ Выявленные риски")
        for r in risks:
            sev = r.get("severity", "")
            emoji = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(sev, "⚪")
            clause = f" `{r['clause']}`" if r.get("clause") else ""
            with st.expander(f"{emoji} {r.get('title', '')} — {r.get('category', '')}{clause}"):
                st.markdown(f"**Описание:** {r.get('description', '')}")
                st.markdown(f"**Рекомендация:** {r.get('recommendation', '')}")

    fz = data.get("fz214", {})
    if fz.get("is_ddu"):
        st.markdown("### 📋 Проверка по 214-ФЗ")
        if fz.get("status") == "COMPLIANT":
            st.success("✅ Соответствует требованиям 214-ФЗ")
        else:
            st.error("❌ Обнаружены нарушения 214-ФЗ")
            for v in fz.get("violations", []):
                st.markdown(f"• {v}")

    st.caption("⚠️ Анализ выполнен ИИ и не является юридической консультацией.")

    # Кнопка скачивания Word (включает правки если они есть)
    st.markdown("---")
    try:
        word_bytes = build_word_report(data, doc_name, corrections)
        file_ts = datetime.now().strftime("%Y%m%d_%H%M")
        label = "📥 Скачать справку с правками (.docx)" if corrections else "📥 Скачать справку в Word (.docx)"
        st.download_button(
            label=label,
            data=word_bytes,
            file_name=f"анализ_договора_{file_ts}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.warning(f"Не удалось сформировать Word-документ: {e}")


def show_corrections(corrections: list[dict]):
    st.markdown("### ✏️ Рекомендуемые правки по рискам высокого уровня")
    for i, c in enumerate(corrections, 1):
        clause = c.get("clause", "")
        title = c.get("risk_title", f"Правка {i}")
        label = f"{i}. {title}" + (f" — `{clause}`" if clause else "")
        with st.expander(label, expanded=True):
            orig = c.get("original_excerpt", "").strip()
            if orig:
                st.markdown("**Текущая редакция:**")
                st.error(orig)
            st.markdown("**Предлагаемая редакция:**")
            st.success(c.get("suggested_text", ""))
            if c.get("rationale"):
                st.caption(f"💡 {c['rationale']}")


# ── UI: основной интерфейс ─────────────────────────────────────────────────

st.set_page_config(page_title="Анализатор договоров", page_icon="📄", layout="wide")
st.title("📄 Анализатор договоров")
st.markdown("Проверка юридических рисков и соответствия **214-ФЗ**")

tab_file, tab_text = st.tabs(["📎 Загрузить файл", "📝 Вставить текст"])

# ── Вкладка: файл ──────────────────────────────────────────────────────────
with tab_file:
    uploaded = st.file_uploader(
        "Выберите файл договора",
        type=["pdf", "docx", "txt"],
        help="Поддерживаются PDF, DOCX, TXT",
    )

    if uploaded and st.button("Анализировать файл", type="primary", key="btn_file"):
        with st.spinner("Читаю документ..."):
            try:
                raw_text = parse_file(uploaded.read(), uploaded.name)
            except Exception as e:
                st.error(f"Ошибка чтения файла: {e}")
                st.stop()

        if len(raw_text.strip()) < 100:
            st.error("Не удалось извлечь текст из документа.")
            st.stop()

        with st.spinner("Анализирую договор (15-30 сек)..."):
            try:
                result = asyncio.run(analyze(raw_text))
            except Exception as e:
                st.error(f"Ошибка анализа: {e}")
                st.stop()

        # Сохраняем в session_state, очищаем старые правки
        st.session_state["result_file"] = result
        st.session_state["text_file"] = raw_text
        st.session_state["name_file"] = uploaded.name
        st.session_state.pop("corrections_file", None)

    # Отображаем результат если он есть
    if "result_file" in st.session_state:
        result = st.session_state["result_file"]
        corrections = st.session_state.get("corrections_file")
        show_results(result, st.session_state.get("name_file", ""), corrections)

        # Кнопка исправления HIGH-рисков
        high_risks = [r for r in result.get("risks", []) if r.get("severity") == "HIGH"]
        if high_risks:
            st.markdown("---")
            if not corrections:
                if st.button(
                    f"✏️ Исправить риски высокого уровня ({len(high_risks)} шт.)",
                    key="btn_correct_file",
                    type="secondary",
                ):
                    with st.spinner("Формирую рекомендации по правкам договора (20-40 сек)..."):
                        try:
                            st.session_state["corrections_file"] = asyncio.run(
                                correct_risks(st.session_state["text_file"], high_risks)
                            )
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ошибка формирования правок: {e}")
            else:
                show_corrections(corrections)
                if st.button("🔄 Пересчитать правки", key="btn_recorrect_file"):
                    st.session_state.pop("corrections_file", None)
                    st.rerun()

# ── Вкладка: текст ─────────────────────────────────────────────────────────
with tab_text:
    contract_text = st.text_area(
        "Вставьте текст договора",
        height=300,
        placeholder="Вставьте полный текст договора...",
    )

    if st.button("Анализировать текст", type="primary", key="btn_text"):
        if len(contract_text.strip()) < 100:
            st.warning("Текст слишком короткий. Вставьте полный текст договора.")
        else:
            with st.spinner("Анализирую договор (15-30 сек)..."):
                try:
                    result = asyncio.run(analyze(contract_text))
                except Exception as e:
                    st.error(f"Ошибка анализа: {e}")
                    st.stop()

            st.session_state["result_text"] = result
            st.session_state["text_text"] = contract_text
            st.session_state.pop("corrections_text", None)

    if "result_text" in st.session_state:
        result = st.session_state["result_text"]
        corrections = st.session_state.get("corrections_text")
        show_results(result, corrections=corrections)

        high_risks = [r for r in result.get("risks", []) if r.get("severity") == "HIGH"]
        if high_risks:
            st.markdown("---")
            if not corrections:
                if st.button(
                    f"✏️ Исправить риски высокого уровня ({len(high_risks)} шт.)",
                    key="btn_correct_text",
                    type="secondary",
                ):
                    with st.spinner("Формирую рекомендации по правкам договора (20-40 сек)..."):
                        try:
                            st.session_state["corrections_text"] = asyncio.run(
                                correct_risks(st.session_state["text_text"], high_risks)
                            )
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ошибка формирования правок: {e}")
            else:
                show_corrections(corrections)
                if st.button("🔄 Пересчитать правки", key="btn_recorrect_text"):
                    st.session_state.pop("corrections_text", None)
                    st.rerun()
