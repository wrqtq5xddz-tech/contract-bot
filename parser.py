import io
import re
from pathlib import Path


def parse_document(content: bytes, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()

    if suffix == ".pdf":
        text = _parse_pdf(content)
    elif suffix == ".docx":
        text = _parse_docx(content)
    elif suffix == ".txt":
        text = _parse_text(content)
    else:
        raise ValueError(f"Неподдерживаемый формат: {suffix}. Поддерживаются: PDF, DOCX, TXT")

    return _clean(text)


def _parse_pdf(content: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _parse_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _parse_text(content: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _clean(text: str) -> str:
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[^\S\n\t]+', ' ', text)
    return text.strip()
